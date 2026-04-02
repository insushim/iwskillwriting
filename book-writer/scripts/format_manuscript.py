#!/usr/bin/env python3
"""
Book Writer - Phase 6: Manuscript Formatter
마크다운 원고를 출판 포맷(DOCX, EPUB, PDF)으로 변환합니다.
"""

import json
import re
import sys
from pathlib import Path
from datetime import datetime


def format_manuscript(project_dir: str, output_format: str = "all") -> dict:
    """원고를 지정된 포맷으로 변환합니다."""
    project_path = Path(project_dir)
    config = _load_config(project_path)
    chapters = _load_chapters(project_path)

    if not chapters:
        return {"error": "챕터 파일이 없습니다."}

    results = {}

    if output_format in ("docx", "all"):
        try:
            docx_path = _create_docx(project_path, config, chapters)
            results["docx"] = {"status": "success", "path": str(docx_path)}
        except ImportError:
            results["docx"] = {"status": "error", "message": "python-docx가 설치되지 않았습니다. pip install python-docx"}
        except Exception as e:
            results["docx"] = {"status": "error", "message": str(e)}

    if output_format in ("epub", "all"):
        try:
            epub_path = _create_epub(project_path, config, chapters)
            results["epub"] = {"status": "success", "path": str(epub_path)}
        except ImportError:
            results["epub"] = {"status": "error", "message": "ebooklib이 설치되지 않았습니다. pip install ebooklib"}
        except Exception as e:
            results["epub"] = {"status": "error", "message": str(e)}

    if output_format in ("pdf", "all"):
        try:
            pdf_path = _create_pdf(project_path, config, chapters)
            results["pdf"] = {"status": "success", "path": str(pdf_path)}
        except ImportError:
            results["pdf"] = {"status": "error", "message": "weasyprint이 설치되지 않았습니다. pip install weasyprint"}
        except Exception as e:
            results["pdf"] = {"status": "error", "message": str(e)}

    # project.json 업데이트
    _update_project(project_path, output_format)

    return results


def _load_config(project_path: Path) -> dict:
    with open(project_path / "project.json", "r", encoding="utf-8") as f:
        return json.load(f)


def _load_chapters(project_path: Path) -> list:
    """챕터를 순서대로 로드합니다."""
    chapters_dir = project_path / "chapters"
    chapters = []

    for ch_file in sorted(chapters_dir.glob("ch*.md")):
        content = ch_file.read_text(encoding="utf-8")
        # 프론트매터 분리
        body = content
        metadata = {}
        if content.startswith("---"):
            parts = content.split("---", 2)
            if len(parts) >= 3:
                body = parts[2].strip()
                for line in parts[1].strip().split("\n"):
                    if ":" in line:
                        key, val = line.split(":", 1)
                        metadata[key.strip()] = val.strip().strip('"')

        # 마크다운 제목 추출
        title_match = re.search(r'^#\s+(.+)$', body, re.MULTILINE)
        title = title_match.group(1) if title_match else metadata.get("title", ch_file.stem)

        # 마크다운 제목 행 제거 (본문만)
        if title_match:
            body = body[title_match.end():].strip()

        chapters.append({
            "num": int(re.search(r'\d+', ch_file.stem).group()),
            "title": title,
            "content": body,
            "metadata": metadata,
        })

    return chapters


def _markdown_to_paragraphs(text: str) -> list:
    """마크다운 텍스트를 단락 리스트로 변환합니다."""
    paragraphs = []
    current_type = "text"

    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue

        if block == "---" or block == "***" or block == "* * *":
            paragraphs.append({"type": "scene_break", "text": "* * *"})
        elif block.startswith("> "):
            paragraphs.append({"type": "blockquote", "text": block[2:]})
        elif block.startswith("## "):
            paragraphs.append({"type": "heading2", "text": block[3:]})
        elif block.startswith("### "):
            paragraphs.append({"type": "heading3", "text": block[4:]})
        elif block.startswith("- ") or block.startswith("* "):
            paragraphs.append({"type": "list_item", "text": block[2:]})
        elif re.match(r'^\d+\.\s', block):
            paragraphs.append({"type": "list_item", "text": re.sub(r'^\d+\.\s', '', block)})
        else:
            # 일반 텍스트 (줄바꿈 처리)
            for line in block.split("\n"):
                line = line.strip()
                if line:
                    paragraphs.append({"type": "text", "text": line})

    return paragraphs


def _create_docx(project_path: Path, config: dict, chapters: list) -> Path:
    """DOCX 파일을 생성합니다."""
    from docx import Document
    from docx.shared import Pt, Mm, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT

    doc = Document()

    # 페이지 설정
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(30)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(30)
    section.right_margin = Mm(25)

    # 기본 스타일
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 1.6
    pf.space_after = Pt(6)

    # 표지 페이지
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(200)
    run = title_para.add_run(config["title"])
    run.font.size = Pt(28)
    run.bold = True

    if config.get("subtitle"):
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle_para.add_run(config["subtitle"])
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(100, 100, 100)

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.paragraph_format.space_before = Pt(50)
    run = author_para.add_run(config.get("author", ""))
    run.font.size = Pt(14)

    doc.add_page_break()

    # 목차 페이지
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("목차")
    run.font.size = Pt(18)
    run.bold = True

    doc.add_paragraph()  # 빈 줄

    for ch in chapters:
        toc_entry = doc.add_paragraph()
        toc_entry.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = toc_entry.add_run(f"{ch['num']}장. {ch['title']}")
        run.font.size = Pt(12)

    doc.add_page_break()

    # 챕터 본문
    for ch in chapters:
        # 챕터 제목
        ch_title = doc.add_paragraph()
        ch_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ch_title.paragraph_format.space_before = Pt(72)
        ch_title.paragraph_format.space_after = Pt(36)
        run = ch_title.add_run(f"{ch['num']}장. {ch['title']}")
        run.font.size = Pt(16)
        run.bold = True

        # 본문
        paragraphs = _markdown_to_paragraphs(ch["content"])
        for para_data in paragraphs:
            if para_data["type"] == "scene_break":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(24)
                run = p.add_run("* * *")
                run.font.size = Pt(12)
            elif para_data["type"] == "blockquote":
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Mm(15)
                run = p.add_run(para_data["text"])
                run.italic = True
                run.font.color.rgb = RGBColor(80, 80, 80)
            elif para_data["type"] in ("heading2", "heading3"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(18)
                run = p.add_run(para_data["text"])
                run.bold = True
                run.font.size = Pt(13 if para_data["type"] == "heading2" else 12)
            else:
                p = doc.add_paragraph()
                # 마크다운 인라인 서식 처리
                text = para_data["text"]
                # Bold
                parts = re.split(r'\*\*(.+?)\*\*', text)
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        if i % 2 == 1:  # 볼드 부분
                            run.bold = True

        # 챕터 끝 페이지 나누기 (마지막 챕터 제외)
        if ch != chapters[-1]:
            doc.add_page_break()

    # 저장
    output_dir = project_path / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    filename = f"{_sanitize(config['title'])}_원고_{datetime.now().strftime('%Y%m%d')}.docx"
    output_path = output_dir / filename
    doc.save(str(output_path))

    return output_path


def _create_epub(project_path: Path, config: dict, chapters: list) -> Path:
    """EPUB 파일을 생성합니다."""
    from ebooklib import epub

    book = epub.EpubBook()

    # 메타데이터
    book.set_identifier(f'book-writer-{_sanitize(config["title"])}')
    book.set_title(config["title"])
    book.set_language(config.get("language", "ko"))
    book.add_author(config.get("author", "Unknown"))

    if config.get("subtitle"):
        book.add_metadata('DC', 'description', config["subtitle"])

    # CSS 스타일
    css_content = """
@charset "UTF-8";
body {
    font-family: serif;
    font-size: 1em;
    line-height: 1.8;
    margin: 1em;
    text-align: justify;
    word-break: keep-all;
}
h1 {
    font-size: 1.5em;
    text-align: center;
    margin-top: 3em;
    margin-bottom: 2em;
    page-break-before: always;
}
p {
    text-indent: 1em;
    margin: 0.3em 0;
}
p.first {
    text-indent: 0;
}
.scene-break {
    text-align: center;
    margin: 2em 0;
    font-size: 1.2em;
}
blockquote {
    margin: 1em 2em;
    font-style: italic;
    color: #555;
}
"""
    css = epub.EpubItem(
        uid="style",
        file_name="style/default.css",
        media_type="text/css",
        content=css_content.encode("utf-8"),
    )
    book.add_item(css)

    # 챕터 생성
    epub_chapters = []
    for ch in chapters:
        c = epub.EpubHtml(
            title=f"{ch['num']}장. {ch['title']}",
            file_name=f"ch{ch['num']:02d}.xhtml",
            lang=config.get("language", "ko"),
        )

        # 마크다운 → HTML 변환
        html_content = _markdown_to_html(ch["content"], ch["num"], ch["title"])
        c.content = html_content.encode("utf-8")
        c.add_item(css)
        book.add_item(c)
        epub_chapters.append(c)

    # 목차
    book.toc = epub_chapters

    # 네비게이션
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    # 스파인
    book.spine = ['nav'] + epub_chapters

    # 저장
    output_dir = project_path / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    filename = f"{_sanitize(config['title'])}_{datetime.now().strftime('%Y%m%d')}.epub"
    output_path = output_dir / filename
    epub.write_epub(str(output_path), book)

    return output_path


def _create_pdf(project_path: Path, config: dict, chapters: list) -> Path:
    """PDF 파일을 생성합니다 (weasyprint 필요)."""
    import weasyprint

    # 전체 HTML 생성
    html_parts = [f"""<!DOCTYPE html>
<html lang="{config.get('language', 'ko')}">
<head>
<meta charset="UTF-8">
<title>{config['title']}</title>
<style>
@page {{
    size: 152mm 225mm;
    margin: 20mm 20mm 25mm 25mm;
    @bottom-center {{
        content: counter(page);
        font-size: 9pt;
    }}
}}
body {{
    font-family: serif;
    font-size: 10pt;
    line-height: 1.7;
    text-align: justify;
}}
h1 {{
    font-size: 16pt;
    text-align: center;
    margin-top: 4em;
    margin-bottom: 2em;
    page-break-before: always;
}}
h1:first-of-type {{
    page-break-before: avoid;
}}
p {{
    text-indent: 1em;
    margin: 0.3em 0;
}}
.scene-break {{
    text-align: center;
    margin: 2em 0;
}}
blockquote {{
    margin: 1em 2em;
    font-style: italic;
    color: #555;
}}
.cover {{
    text-align: center;
    margin-top: 40%;
    page-break-after: always;
}}
.cover h1 {{
    font-size: 28pt;
    page-break-before: avoid;
}}
</style>
</head>
<body>
<div class="cover">
<h1>{config['title']}</h1>
"""]

    if config.get("subtitle"):
        html_parts.append(f"<p style='font-size: 14pt; color: #666;'>{config['subtitle']}</p>")

    html_parts.append(f"<p style='font-size: 12pt; margin-top: 3em;'>{config.get('author', '')}</p>")
    html_parts.append("</div>")

    for ch in chapters:
        html_parts.append(_markdown_to_html(ch["content"], ch["num"], ch["title"]))

    html_parts.append("</body></html>")
    full_html = "\n".join(html_parts)

    output_dir = project_path / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    filename = f"{_sanitize(config['title'])}_인쇄용_{datetime.now().strftime('%Y%m%d')}.pdf"
    output_path = output_dir / filename

    html = weasyprint.HTML(string=full_html)
    html.write_pdf(str(output_path))

    return output_path


def _markdown_to_html(content: str, chapter_num: int, chapter_title: str) -> str:
    """마크다운을 HTML로 변환합니다."""
    html = f"<h1>{chapter_num}장. {chapter_title}</h1>\n"

    paragraphs = _markdown_to_paragraphs(content)
    is_first = True

    for para in paragraphs:
        if para["type"] == "scene_break":
            html += '<p class="scene-break">* * *</p>\n'
            is_first = True
        elif para["type"] == "blockquote":
            html += f'<blockquote>{_escape_html(para["text"])}</blockquote>\n'
        elif para["type"] == "heading2":
            html += f'<h2>{_escape_html(para["text"])}</h2>\n'
            is_first = True
        elif para["type"] == "heading3":
            html += f'<h3>{_escape_html(para["text"])}</h3>\n'
        elif para["type"] == "list_item":
            html += f'<li>{_escape_html(para["text"])}</li>\n'
        else:
            text = _escape_html(para["text"])
            # Bold
            text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
            # Italic
            text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)

            css_class = ' class="first"' if is_first else ''
            html += f'<p{css_class}>{text}</p>\n'
            is_first = False

    return html


def _escape_html(text: str) -> str:
    """HTML 특수문자를 이스케이프합니다."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _sanitize(name: str) -> str:
    """파일명에 사용할 수 없는 문자를 제거합니다."""
    invalid = '<>:"/\\|?*'
    for ch in invalid:
        name = name.replace(ch, "")
    return name.strip().replace(" ", "_")


def _update_project(project_path: Path, output_format: str):
    """프로젝트 상태를 업데이트합니다."""
    config_path = project_path / "project.json"
    with open(config_path, "r", encoding="utf-8") as f:
        config = json.load(f)

    config["updated_at"] = datetime.now().isoformat()
    config["current_phase"] = max(config.get("current_phase", 0), 6)
    config["status"] = "formatted"

    with open(config_path, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python format_manuscript.py <project_dir> [--format docx|epub|pdf|all]")
        sys.exit(1)

    project_dir = sys.argv[1]
    fmt = "all"
    if "--format" in sys.argv:
        idx = sys.argv.index("--format")
        if idx + 1 < len(sys.argv):
            fmt = sys.argv[idx + 1]

    result = format_manuscript(project_dir, fmt)
    print(json.dumps(result, ensure_ascii=False, indent=2))
